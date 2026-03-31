"""
Generates CX_Opportunities_Dashboard_Plan_2026.docx from structured content.
Uses python-docx to produce a professional Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shading_elm)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, "2E4057")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = str(val)
            for p in row.cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CX Opportunities Dashboard\n2026 Plan")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Draft for Review — March 31, 2026")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("CX Analytics & Strategy Team")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── Document Control ──
    doc.add_heading("Document Control", level=1)
    add_table(doc,
        ["Field", "Value"],
        [
            ["Version", "1.0 — Draft"],
            ["Date", "March 31, 2026"],
            ["Author", "CX Analytics & Strategy Team"],
            ["Status", "Draft for Review"],
            ["Distribution", "CX Leadership, Product, Technology, BI"],
        ],
        col_widths=[5, 12],
    )
    doc.add_paragraph("")

    # ── 1. Executive Summary ──
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document outlines the plan for a Looker-based analytics dashboard that "
        "tracks the 7 highest-impact improvement opportunities identified by the CX "
        "organization for 2026. Each opportunity maps directly to one or more of the "
        "5 CX Goals for the year, allowing leadership to monitor whether day-to-day "
        "execution against these opportunities is translating into progress on the "
        "annual targets."
    )
    doc.add_paragraph("The dashboard suite will consist of:")
    add_table(doc,
        ["Dashboard", "Purpose"],
        [
            ["CX Executive KPI Overview", "Bird's-eye view of all 5 goals, their KPIs, and YTD progress"],
            ["Opportunity Impact Tracker", "Status and projected impact of each of the 7 opportunities"],
            ["Opportunity Deep-Dive (x7)", "One drill-down dashboard per opportunity with operational metrics"],
        ],
        col_widths=[6, 12],
    )
    doc.add_paragraph("")

    # ── 2. The 5 CX Goals ──
    doc.add_heading("2. The 5 CX Goals for 2026", level=1)

    goals = [
        {
            "name": "Goal 1 — Faster, More Effective 1:1 Support",
            "desc": "Provide faster, more effective 1:1 support across all CX touchpoints.",
            "kpis": [
                ["First Contact Resolution", "≥ 90%", "TBD — pull from existing FCR reporting"],
                ["Total Time per Contact", "Reduce by 2 minutes", "TBD — pull from AHT reporting"],
            ],
        },
        {
            "name": "Goal 2 — Automate Where Humans Don't Add Value",
            "desc": "Automate customer interactions where a human doesn't add value.",
            "kpis": [
                ["Self-Service Rate", "≥ 75% (+5 pp with automation)", "TBD"],
                ["Chatbot CSAT", "+20 pp improvement", "TBD"],
                ["Advocate Accuracy Rate", "90% within 30 days of training/event/release", "TBD"],
            ],
        },
        {
            "name": "Goal 3 — AI & Automation for Quality at Scale",
            "desc": "Leverage AI & Automation to reinforce quality at scale.",
            "kpis": [
                ["Global CSAT", "≥ 94%", "TBD"],
                ["Market-to-Market CSAT Variance", "< 5 points", "TBD"],
            ],
        },
        {
            "name": "Goal 4 — Optimize Top Cost Drivers",
            "desc": "Aggressively optimize top cost drivers to drive efficiency, quality, and scale.",
            "kpis": [
                ["Global Cost per Case (YoY)", "15% reduction by EOY", "TBD"],
                ["Run-Rate Savings", "$$$M delivered by EOY", "TBD"],
            ],
        },
        {
            "name": "Goal 5 — Eliminate Top Contact Drivers",
            "desc": "Identify and eliminate top contact drivers to improve growth, satisfaction, and engagement.",
            "kpis": [
                ["Contact Rate", "25% reduction", "TBD"],
                ["Improvements Surfaced", "≥ 15 improvements driving engagement/growth", "TBD"],
                ["Tech Contact Rate", "≤ 4.0", "TBD"],
            ],
        },
    ]

    for g in goals:
        doc.add_heading(g["name"], level=2)
        doc.add_paragraph(g["desc"])
        add_table(doc,
            ["KPI", "2026 Target", "Baseline (EOY 2025)"],
            g["kpis"],
            col_widths=[5, 6, 6],
        )
        doc.add_paragraph("")

    # ── 3. The 7 Opportunities ──
    doc.add_heading("3. The 7 Opportunities — Descriptions & Goal Mapping", level=1)
    doc.add_paragraph(
        "Each opportunity is mapped to the goals it most directly supports, along "
        "with the primary metrics it is expected to move."
    )

    opps = [
        {
            "id": "OPP-1",
            "name": "Reducing Incorrect / Inaccessible Email Address Contacts",
            "desc": (
                "A significant volume of inbound contacts stems from customers unable "
                "to access their accounts due to incorrect or outdated email addresses. "
                "Reducing these contacts requires both proactive email hygiene initiatives "
                "and improved self-service flows for email updates."
            ),
            "goals": "Goal 1 (1:1 Support), Goal 5 (Eliminate Contact Drivers)",
            "metrics": "Contact volume from email issues, FCR on email tickets, self-service completion rate for email changes",
            "sources": "CRM / ticketing system (contact reason taxonomy), IVR/chatbot logs, account management funnel data",
        },
        {
            "id": "OPP-2",
            "name": "Account Sharing Inquiry Surge — Mitigation with Product & Technology",
            "desc": (
                "Policy and product changes around account sharing are expected to "
                "generate a spike in customer inquiries. Proactive collaboration with "
                "Product and Technology is needed to deflect avoidable contacts through "
                "better in-app messaging, FAQ content, and policy clarity."
            ),
            "goals": "Goal 2 (Automation), Goal 4 (Cost Optimization), Goal 5 (Eliminate Contact Drivers)",
            "metrics": "Account sharing contact volume, deflection rate, chatbot resolution rate for sharing topics, CSAT on sharing interactions",
            "sources": "CRM / ticketing (contact reason), chatbot topic classification, product telemetry (sharing feature usage), CSAT surveys",
        },
        {
            "id": "OPP-3",
            "name": "Understanding OTP Journey Pain Points",
            "desc": (
                "One-Time Password (OTP) flows are a friction point in the authentication "
                "experience. Deeper data exploration is required to identify where customers "
                "are failing, retrying, or abandoning the OTP journey."
            ),
            "goals": "Goal 1 (1:1 Support), Goal 5 (Eliminate Contact Drivers)",
            "metrics": "OTP-related contact volume, OTP success/failure rate by step, OTP retry rate, time to complete OTP, OTP-triggered churn",
            "sources": "Authentication service logs, CRM (contact reason), product analytics (funnel events), chatbot logs",
        },
        {
            "id": "OPP-4",
            "name": "Disney Bundle Activation — Customer Confusion",
            "desc": (
                "Customers are confused during the Disney bundle activation flow, resulting "
                "in contacts and negative sentiment. Root-cause analysis of where confusion "
                "arises (eligibility, linking, billing) is needed to drive UX and communication fixes."
            ),
            "goals": "Goal 1 (1:1 Support), Goal 3 (Quality at Scale), Goal 5 (Eliminate Contact Drivers)",
            "metrics": "Bundle activation contact volume, activation success rate, CSAT on bundle interactions, repeat contact rate for bundle issues",
            "sources": "CRM (contact reason), bundle activation funnel data, CSAT surveys, chatbot/IVR topic logs",
        },
        {
            "id": "OPP-5",
            "name": "Automation for FCCA & Third-Party Cancellations",
            "desc": (
                "First Contact Cancellation Attempts (FCCA) and third-party cancellations "
                "represent repeatable, rules-based interactions that are strong candidates "
                "for automation. Evaluating and implementing automation will reduce cost "
                "and handle time while maintaining quality."
            ),
            "goals": "Goal 2 (Automation), Goal 4 (Cost Optimization)",
            "metrics": "FCCA volume, third-party cancellation volume, automation rate, AHT for cancellation contacts, cost per cancellation case, CSAT post-cancellation",
            "sources": "CRM / ticketing, billing system (cancellation events), workforce management (AHT), cost accounting, CSAT surveys",
        },
        {
            "id": "OPP-6",
            "name": "Refund Issues & Automation Opportunities",
            "desc": (
                "Refund requests are a high-volume contact driver. Identifying which refund "
                "scenarios can be safely automated — and which require human judgment — will "
                "reduce volume, cost, and customer effort."
            ),
            "goals": "Goal 2 (Automation), Goal 4 (Cost Optimization), Goal 5 (Eliminate Contact Drivers)",
            "metrics": "Refund contact volume, refund automation rate, refund processing time, refund CSAT, cost per refund case",
            "sources": "CRM / ticketing, billing/payment system, automation platform logs, CSAT surveys",
        },
        {
            "id": "OPP-7",
            "name": "Additional Deflection Levers to Reduce Spam",
            "desc": (
                "Spam and non-actionable contacts waste agent capacity. Identifying and "
                "implementing additional deflection mechanisms (better IVR trees, CAPTCHA, "
                "authentication gates) will free capacity for genuine customer needs."
            ),
            "goals": "Goal 4 (Cost Optimization), Goal 5 (Eliminate Contact Drivers)",
            "metrics": "Spam contact volume, spam deflection rate, false-positive rate, agent utilization gained",
            "sources": "CRM / ticketing (spam classification), IVR/chatbot logs, telephony platform data",
        },
    ]

    for opp in opps:
        doc.add_heading(f"{opp['id']}: {opp['name']}", level=2)
        doc.add_paragraph(opp["desc"])
        add_table(doc,
            ["Attribute", "Detail"],
            [
                ["Primary Goals", opp["goals"]],
                ["Key Metrics", opp["metrics"]],
                ["Data Sources", opp["sources"]],
            ],
            col_widths=[4, 13],
        )
        doc.add_paragraph("")

    # ── 4. Opportunity → Goal Mapping Matrix ──
    doc.add_heading("4. Opportunity → Goal Mapping Matrix", level=1)
    BULLET = "●"
    mapping_rows = [
        ["OPP-1: Email Addresses",  BULLET, "",     "",     "",     BULLET],
        ["OPP-2: Account Sharing",  "",     BULLET, "",     BULLET, BULLET],
        ["OPP-3: OTP Journey",      BULLET, "",     "",     "",     BULLET],
        ["OPP-4: Disney Bundle",    BULLET, "",     BULLET, "",     BULLET],
        ["OPP-5: FCCA / 3P Cancel", "",     BULLET, "",     BULLET, ""],
        ["OPP-6: Refunds",          "",     BULLET, "",     BULLET, BULLET],
        ["OPP-7: Spam Deflection",  "",     "",     "",     BULLET, BULLET],
    ]
    t = add_table(doc,
        ["Opportunity", "Goal 1:\n1:1 Support", "Goal 2:\nAutomation", "Goal 3:\nQuality", "Goal 4:\nCost Opt.", "Goal 5:\nContact Drivers"],
        mapping_rows,
        col_widths=[4.5, 2.5, 2.5, 2.5, 2.5, 3],
    )
    for row in t.rows[1:]:
        for cell in row.cells[1:]:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── 5. Dashboard Architecture ──
    doc.add_heading("5. Dashboard Architecture", level=1)

    doc.add_heading("5.1  Dashboard 1 — CX Executive KPI Overview", level=2)
    doc.add_paragraph("Audience: CX VP, Directors, cross-functional leadership")
    doc.add_paragraph(
        "This dashboard provides a bird's-eye view of the 5 CX goals. "
        "It features color-coded KPI scorecards at the top, a multi-line time series "
        "tracking each goal's primary metric over time, a waterfall chart showing the "
        "estimated contact-volume reduction contributed by each opportunity, and a market "
        "CSAT heatmap."
    )
    add_table(doc,
        ["Tile", "Visualization", "Description"],
        [
            ["Goal KPI Scorecards (x5)", "Single-value tiles", "Current value, target, and RAG status for each goal's primary KPI"],
            ["Goal Progress Over Time", "Multi-line time series", "Weekly/monthly trend for FCR, Self-Service %, CSAT %, Cost per Case, Contact Rate"],
            ["Opportunity Impact Waterfall", "Waterfall chart", "Estimated contact-volume reduction contributed by each of the 7 opportunities"],
            ["Goal Health Heatmap", "Heatmap", "Rows = markets, columns = goals; cell color = RAG status (green/amber/red)"],
        ],
        col_widths=[5, 4, 8],
    )
    doc.add_paragraph("")

    doc.add_heading("5.2  Dashboard 2 — Opportunity Impact Tracker", level=2)
    doc.add_paragraph("Audience: CX Directors, Opportunity Owners, Program Managers")
    doc.add_paragraph(
        "This dashboard tracks the status and measurable impact of each of the 7 "
        "opportunities. It features a status table with RAG indicators, monthly contact "
        "volume by opportunity, projected vs actual cost savings, and a timeline view."
    )
    add_table(doc,
        ["Tile", "Visualization", "Description"],
        [
            ["Opportunity Status Table", "Table with conditional formatting", "One row per opportunity: RAG status, owner, estimated vs. actual impact, gap"],
            ["Contact Volume by Opportunity", "Stacked bar chart (monthly)", "Shows contact volume attributable to each opportunity over time"],
            ["Projected vs Actual Savings", "Dual-axis line/bar", "Compares forecasted cost/volume savings against actuals"],
            ["Goal Contribution", "Sankey / alluvial diagram", "Visual flow from opportunities to goals showing contribution relationships"],
            ["Opportunity Timeline", "Gantt chart", "Milestones and phases for each opportunity's implementation"],
            ["Cumulative YTD Impact", "Stacked area chart", "Running total of contact volume reduction across all opportunities"],
        ],
        col_widths=[5, 4, 8],
    )
    doc.add_paragraph("")

    doc.add_heading("5.3  Dashboard 3 (x7) — Opportunity Deep-Dives", level=2)
    doc.add_paragraph(
        "Each of the 7 opportunities gets its own drill-down dashboard. The layout is "
        "consistent across all 7 (4 KPI scorecards, a time-series trend, a breakdown "
        "chart, and a drillable detail table), but the specific metrics vary:"
    )
    add_table(doc,
        ["Opportunity", "Scorecard Metrics", "Breakdown Dimension", "Root Cause View"],
        [
            ["OPP-1: Email", "Email-contact vol, FCR, self-service rate, repeat rate", "By market, by email issue sub-type", "Top reasons for email inaccessibility"],
            ["OPP-2: Account Sharing", "Sharing vol, deflection rate, bot resolution, CSAT", "By market, by inquiry sub-type", "Sharing policy confusion categories"],
            ["OPP-3: OTP Journey", "OTP vol, success rate, retry rate, abandonment rate", "By OTP step, by device/platform", "Funnel drop-off analysis"],
            ["OPP-4: Disney Bundle", "Bundle vol, activation success rate, CSAT, repeat rate", "By confusion category, by channel", "Activation funnel failure points"],
            ["OPP-5: FCCA / 3P Cancel", "Cancel vol, automation rate, AHT, cost per case", "By cancellation type, by channel", "Automation eligibility segmentation"],
            ["OPP-6: Refunds", "Refund vol, automation rate, processing time, CSAT", "By refund reason, by amount tier", "Automation vs. human-required segments"],
            ["OPP-7: Spam Deflection", "Spam vol, deflection rate, false-positive rate, capacity freed", "By spam type, by entry channel", "Deflection mechanism effectiveness"],
        ],
        col_widths=[3.5, 5, 4.5, 4.5],
    )
    doc.add_paragraph("")

    # ── 6. Data Requirements ──
    doc.add_heading("6. Data Requirements & Sources", level=1)

    doc.add_heading("6.1  Required Data Sources", level=2)
    add_table(doc,
        ["Source", "Key Tables / Feeds", "Refresh", "Owner"],
        [
            ["CRM / Ticketing System", "Cases, contact reasons, dispositions", "Daily", "CX Ops"],
            ["CSAT Survey Platform", "Survey responses, scores, verbatims", "Daily", "CX Analytics"],
            ["Chatbot / IVR Platform", "Conversations, intents, outcomes", "Daily", "Digital CX"],
            ["Billing / Payment System", "Transactions, refunds, cancellations", "Daily", "Finance / Billing"],
            ["Authentication Service", "OTP events, success/fail/retry logs", "Daily", "Technology"],
            ["Product Analytics", "Funnel events, feature usage", "Daily", "Product Analytics"],
            ["Workforce Management", "AHT, schedule adherence, utilization", "Daily", "WFM"],
            ["Cost Accounting", "Cost per case, headcount, vendor costs", "Monthly", "Finance"],
            ["Bundle Activation Funnel", "Activation attempts, completions", "Daily", "Product"],
            ["Spam Classification", "Spam flags, confidence scores", "Daily", "CX Ops"],
        ],
        col_widths=[4.5, 6, 2, 4],
    )
    doc.add_paragraph("")

    doc.add_heading("6.2  Data Model Summary", level=2)
    doc.add_paragraph("The LookML model is organized around these primary views:")
    data_model_items = [
        ("cx_cases", "Core fact table — one row per customer contact/case"),
        ("cx_csat_surveys", "CSAT survey responses joined to cases"),
        ("cx_goals", "Goal definitions, KPIs, and targets (dimension table)"),
        ("cx_opportunities", "Opportunity definitions and metadata (dimension table)"),
        ("cx_opportunity_metrics", "Time-series metrics for each opportunity"),
        ("cx_goal_kpis", "Time-series KPI values for each goal"),
    ]
    for name, desc in data_model_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(name)
        run.bold = True
        p.add_run(f" — {desc}")
    doc.add_paragraph("")

    # ── 7. Implementation Phases ──
    doc.add_heading("7. Implementation Phases", level=1)

    phases = [
        ("Phase 1 — Foundation", [
            "Finalize data source access and ETL pipelines",
            "Build LookML views for core data (cases, CSAT, goals, opportunities)",
            "Deploy CX Executive KPI Overview dashboard",
            "Validate KPI calculations with CX Analytics team",
        ]),
        ("Phase 2 — Opportunity Tracking", [
            "Build opportunity_metrics data pipeline",
            "Deploy Opportunity Impact Tracker dashboard",
            "Set up automated alerting for KPI threshold breaches",
            "Validate opportunity-to-goal attribution logic",
        ]),
        ("Phase 3 — Deep-Dives", [
            "Deploy deep-dive dashboards for OPP-1 through OPP-7",
            "Integrate additional data sources (OTP logs, bundle funnel, spam classification)",
            "Build drill-through links between dashboards",
            "User acceptance testing with opportunity owners",
        ]),
        ("Phase 4 — Optimization (Ongoing)", [
            "Refine visualizations based on stakeholder feedback",
            "Add predictive/forecasting tiles where data supports it",
            "Automate monthly executive summary delivery via Looker schedules",
            "Iterate on goal-to-opportunity attribution as new data becomes available",
        ]),
    ]
    for phase_name, items in phases:
        doc.add_heading(phase_name, level=2)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("")

    # ── 8. Access & Governance ──
    doc.add_heading("8. Access & Governance", level=1)
    add_table(doc,
        ["Role", "Access Level"],
        [
            ["CX VP / SVP", "Full access to all dashboards"],
            ["CX Directors", "Full access to all dashboards"],
            ["Opportunity Owners", "Full access to their opportunity deep-dive + tracker"],
            ["CX Managers", "Read access to Executive KPI + Tracker"],
            ["Cross-Functional (Product, Tech)", "Read access to relevant deep-dives"],
        ],
        col_widths=[6, 11],
    )
    doc.add_paragraph("")

    # ── 9. Open Questions & Risks ──
    doc.add_heading("9. Open Questions & Risks", level=1)
    add_table(doc,
        ["#", "Item", "Status"],
        [
            ["1", "Baseline values for all KPIs need to be confirmed from EOY 2025 reporting", "Open"],
            ["2", "Cost per case data granularity — available at case level or only aggregated?", "Open"],
            ["3", "OTP authentication logs — access and format TBD with Technology team", "Open"],
            ["4", "Account Sharing contact volume projections — dependent on Product roadmap timing", "Open"],
            ["5", "Spam classification model accuracy — validate false-positive rates before dashboard use", "Open"],
            ["6", "Run-rate savings target ($$$M) — exact figure to be confirmed by Finance", "Open"],
        ],
        col_widths=[1, 13, 2.5],
    )
    doc.add_paragraph("")

    # ── 10. Appendix — Looker Import Guide ──
    doc.add_heading("10. Appendix — How to Import Into Looker", level=1)

    doc.add_paragraph(
        "The LookML files in this repository are ready to import into a Looker instance. "
        "Below is a step-by-step guide for getting the dashboards live."
    )

    doc.add_heading("Option A — Connect via Git (Recommended)", level=2)
    steps_a = [
        "In Looker, go to Develop → Manage LookML Projects → New LookML Project.",
        "Choose \"Clone Public Git Repository\" and paste this repository's URL.",
        "Open the project. Looker will automatically discover the .model.lkml, .view.lkml, "
        ".explore.lkml, and .dashboard.lookml files.",
        "Open cx_opportunities.model.lkml and change the connection: line to your "
        "database connection name (the one configured in Admin → Connections).",
        "Update the sql_table_name in each view file to match your actual schema and table names.",
        "Click \"Validate LookML\" to check for errors.",
        "Deploy to production when ready.",
    ]
    for i, step in enumerate(steps_a, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("Option B — Copy-Paste Into an Existing Project", level=2)
    steps_b = [
        "In Looker, open your existing LookML project in Development Mode (Develop → your project).",
        "Create a new model file and paste the contents of lookml/cx_opportunities.model.lkml.",
        "Create new view files under a views/ folder and paste each .view.lkml file's contents.",
        "Create explore and dashboard files the same way.",
        "Update the connection: and sql_table_name references to match your environment.",
        "Validate and deploy.",
    ]
    for i, step in enumerate(steps_b, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("Option C — Use Looker API (for CI/CD pipelines)", level=2)
    doc.add_paragraph(
        "If your organization uses a CI/CD pipeline for LookML deployments, you can push "
        "this repository's lookml/ directory to your LookML project's Git remote. Looker "
        "will automatically pick up the files on the next deploy cycle."
    )

    doc.add_heading("What You'll Need Before Importing", level=2)
    prereqs = [
        "A Looker instance with Developer or Admin access",
        "A database connection configured in Admin → Connections that points to your data warehouse",
        "The actual schema and table names for your CX data (the view files use placeholder names like schema.cx_cases)",
        "The seed data SQL scripts run against your warehouse (lookml/seed_data/) to populate the goals and opportunities dimension tables",
    ]
    for item in prereqs:
        doc.add_paragraph(item, style="List Bullet")

    # ── Save ──
    out_path = "/workspace/docs/CX_Opportunities_Dashboard_Plan_2026.docx"
    doc.save(out_path)
    print(f"Word document saved to {out_path}")


if __name__ == "__main__":
    build_document()
